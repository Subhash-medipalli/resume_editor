"""Clone the parent Word resume and apply structural edits, keeping layout.

Edits are planned by aligning the base resume's blocks against the tailored
resume's blocks in document order, so a paragraph can only ever be matched with
something near it. Additions and deletions are applied to the document rather
than dropped, and paragraph rewrites reuse the original runs so intra-line
formatting (a bold "Languages & Frameworks:" label, say) survives.
"""

from __future__ import annotations

import copy
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from resume_tailor.structure import (
    Block,
    from_markdown,
    parse_paragraphs,
    to_markdown,
)

# Below this similarity two blocks are treated as unrelated rather than an edit.
MIN_REWRITE_RATIO = 0.35


def find_parent_docx(root: Path) -> Path | None:
    resume_dir = root / "resume"
    preferred = resume_dir / "Sravya_M_resume.docx"
    if preferred.is_file():
        return preferred
    if not resume_dir.is_dir():
        return None
    for path in sorted(resume_dir.glob("*.docx")):
        name = path.name.lower()
        if "tailored" in name or name.endswith(".original.docx"):
            continue
        return path
    return None


def iter_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def nonempty_paragraphs(document: Document) -> list[Paragraph]:
    return [p for p in iter_paragraphs(document) if (p.text or "").strip()]


def document_blocks(document: Document) -> tuple[list[Block], list[Paragraph]]:
    paragraphs = nonempty_paragraphs(document)
    entries = [(p.text, p.style.name) for p in paragraphs]
    return parse_paragraphs(entries), paragraphs


def extract_blocks(path: Path) -> list[Block]:
    blocks, _ = document_blocks(Document(str(path)))
    return blocks


def extract_markdown(path: Path) -> str:
    """The structured markdown handed to the model and to the guardrails."""
    return to_markdown(extract_blocks(path))


def extract_text(path: Path) -> str:
    """Flat text of the document. Kept for callers that only need the words."""
    document = Document(str(path))
    return "\n".join(p.text for p in nonempty_paragraphs(document))


def _has_hyperlink(paragraph: Paragraph) -> bool:
    """Hyperlink text lives outside paragraph.runs, so rewriting duplicates it."""
    links = getattr(paragraph, "hyperlinks", None)
    if links:
        return True
    return paragraph._p.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
    ) != []


# Models routinely swap curly quotes for straight ones (and vice versa). That is
# typography, not a tailoring edit, and rewriting the paragraph for it would
# churn the document and lose run formatting for nothing.
_TYPOGRAPHY = str.maketrans({
    "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "-", "\u00a0": " ",
})


def _typographically_equal(a: str, b: str) -> bool:
    return a.translate(_TYPOGRAPHY).strip() == b.translate(_TYPOGRAPHY).strip()


def _assign_label_value(runs, text: str) -> bool:
    """Lay "Label: value" across a bold label run and a regular value run.

    Needed whenever the new text shares no prefix with the old: without it the
    whole string lands in runs[0] and inherits the label's bold, turning an
    entire skills line bold.
    """
    marker = text.find(": ")
    if not runs or not runs[0].bold or not 0 < marker < 60:
        return False
    value_index = next((i for i, run in enumerate(runs) if not run.bold), None)
    if not value_index:
        return False
    runs[0].text = text[: marker + 2]
    runs[value_index].text = text[marker + 2 :]
    for index, run in enumerate(runs):
        if index not in (0, value_index):
            run.text = ""
    return True


def _common_prefix_len(a: str, b: str) -> int:
    limit = min(len(a), len(b))
    i = 0
    while i < limit and a[i] == b[i]:
        i += 1
    return i


def set_paragraph_text(paragraph: Paragraph, text: str) -> bool:
    """Rewrite a paragraph, preserving as much run formatting as possible.

    Runs lying entirely inside the unchanged prefix are left alone; the first run
    that reaches past it absorbs the remainder. Returns False when nothing was
    changed.
    """
    current = paragraph.text or ""
    if _typographically_equal(current, text):
        return False  # identical but for whitespace or quote style
    if _has_hyperlink(paragraph):
        return False  # rewriting would duplicate the linked text

    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text.strip())
        return True

    old = "".join(run.text for run in runs)
    leading = old[: len(old) - len(old.lstrip())]
    new = text.strip()
    if leading:
        new = leading + new  # keep space-based centring

    keep = _common_prefix_len(old, new)
    # If the change reaches into the bold label run, letting that run absorb the
    # remainder would embolden the rest of the line.
    if keep < len(runs[0].text) and runs[0].bold and _assign_label_value(runs, new):
        return True
    consumed = 0
    absorbed = False
    for run in runs:
        start = consumed
        consumed += len(run.text)
        if consumed <= keep and not absorbed:
            continue  # wholly inside the unchanged prefix
        if not absorbed:
            run.text = new[start:]
            absorbed = True
        else:
            run.text = ""
    if not absorbed:
        runs[-1].text = runs[-1].text + new[keep:]
    return True


def _clone_paragraph_after(anchor: Paragraph, template: Paragraph, text: str) -> Paragraph:
    """Insert a new paragraph after `anchor`, formatted like `template`."""
    element = copy.deepcopy(template._p)
    # Keep the template's run structure (a bold label followed by regular text
    # is meaningful); the text of every run is reassigned or blanked below.
    for link in element.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
    ):
        link.getparent().remove(link)
    anchor._p.addnext(element)
    created = Paragraph(element, anchor._parent)
    new_text = text.strip()
    if not created.runs:
        created.add_run(new_text)
        return created

    # A "Label: value" line in this document is a bold label followed by regular
    # text. Dropping the whole string into runs[0] would inherit the label's bold
    # and embolden the entire line, so split it the way the template is split.
    if not _assign_label_value(created.runs, new_text):
        created.runs[0].text = new_text
        for run in created.runs[1:]:
            run.text = ""
    return created


def _delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def plan_edits(base: list[Block], tailored: list[Block]) -> list[tuple]:
    """Align base and tailored blocks in order; return ('replace'|'insert'|'delete', ...).

    Ordered alignment means a paragraph is only ever paired with a nearby one, so
    bullets cannot swap between employers the way global best-match allowed.
    """
    matcher = SequenceMatcher(
        a=[block.key for block in base],
        b=[block.key for block in tailored],
    )
    operations: list[tuple] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag == "replace":
            overlap = min(i2 - i1, j2 - j1)
            for offset in range(overlap):
                old, new = base[i1 + offset], tailored[j1 + offset]
                ratio = SequenceMatcher(None, old.text, new.text).ratio()
                if ratio >= MIN_REWRITE_RATIO or old.kind == new.kind:
                    operations.append(("replace", old, new))
                else:
                    # Insert first: anchoring on a paragraph that has already
                    # been detached puts the clone outside the document, losing
                    # both the old line and its replacement.
                    operations.append(("insert", old, new))
                    operations.append(("delete", old, None))
            for extra in range(i1 + overlap, i2):
                operations.append(("delete", base[extra], None))
            for extra in range(j1 + overlap, j2):
                anchor = base[i1 + overlap - 1] if overlap else base[max(i1 - 1, 0)]
                operations.append(("insert", anchor, tailored[extra]))
        elif tag == "delete":
            for index in range(i1, i2):
                operations.append(("delete", base[index], None))
        elif tag == "insert":
            anchor = base[i1 - 1] if i1 > 0 else None
            for offset in range(j1, j2):
                operations.append(("insert", anchor, tailored[offset]))
    return operations


def _template_for(
    kind: str,
    near: int,
    base_blocks: list[Block],
    paragraphs: list[Paragraph],
) -> Paragraph | None:
    """The nearest existing paragraph of the same kind, to copy formatting from.

    Cloning the anchor is wrong when the anchor is a section heading: an inserted
    skills line would inherit the heading's all-bold runs and come out entirely
    bold. A sibling of the same kind carries the right formatting.
    """
    candidates = [
        index for index, block in enumerate(base_blocks) if block.kind == kind
    ]
    if not candidates:
        return None
    closest = min(candidates, key=lambda index: abs(index - near))
    para_index = base_blocks[closest].para_index
    if para_index is None:
        return None
    return paragraphs[para_index]


def verify_written_docx(dest: Path, approved_markdown: str) -> list[str]:
    """Compare the saved .docx against the text the guardrails approved.

    Nothing else in the pipeline checks that the document on disk matches what
    was reviewed, which is exactly how a silent content loss reaches an employer.
    """
    def normalise(blocks):
        # Compare on the same identity used to align them: whitespace, trailing
        # colons and quote style are presentation, not content.
        return [
            (block.key[1].translate(_TYPOGRAPHY), block.text)
            for block in blocks
            if block.text.strip()
        ]

    written = normalise(extract_blocks(dest))
    approved = normalise(from_markdown(approved_markdown))
    if [key for key, _ in written] == [key for key, _ in approved]:
        return []

    written_keys = {key for key, _ in written}
    approved_keys = {key for key, _ in approved}
    problems: list[str] = []
    for key, text in approved:
        if key not in written_keys:
            problems.append(f"Approved line is missing from the Word file: {text[:90]}")
    for key, text in written:
        if key not in approved_keys:
            problems.append(f"Word file contains an unapproved line: {text[:90]}")
    if not problems:
        problems.append(
            "The Word file's paragraph order does not match the approved resume."
        )
    return problems[:6]


def _surviving_anchor(
    index: int, removed: set[int], paragraphs: list[Paragraph]
) -> Paragraph | None:
    """The nearest paragraph at or before `index` that is still in the document."""
    for candidate in range(index, -1, -1):
        if candidate not in removed:
            return paragraphs[candidate]
    return None


def write_tailored_docx(parent: Path, tailored_text: str, dest: Path) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(parent))
    base_blocks, paragraphs = document_blocks(document)
    tailored_blocks = from_markdown(tailored_text)
    if not tailored_blocks:
        document.save(str(dest))
        return dest

    removed: set[int] = set()
    # Where the next insert for a given anchor should go. Without this every
    # clone is placed immediately after the same anchor, so a run of added
    # lines comes out in reverse order.
    cursor: dict[int, Paragraph] = {}

    for kind, anchor_block, new_block in plan_edits(base_blocks, tailored_blocks):
        if anchor_block is None or anchor_block.para_index is None:
            continue
        index = anchor_block.para_index

        if kind == "replace":
            if index not in removed:
                set_paragraph_text(paragraphs[index], new_block.text)
        elif kind == "delete":
            if index not in removed:
                _delete_paragraph(paragraphs[index])
                removed.add(index)
        elif kind == "insert":
            anchor = cursor.get(index)
            if anchor is None:
                anchor = _surviving_anchor(index, removed, paragraphs)
            if anchor is None:
                continue
            template = _template_for(
                new_block.kind, index, base_blocks, paragraphs
            ) or anchor
            cursor[index] = _clone_paragraph_after(anchor, template, new_block.text)

    document.save(str(dest))
    return dest
