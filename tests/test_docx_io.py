"""Word output fidelity: the file the user actually sends to an employer."""

import pytest

docx = pytest.importorskip("docx")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from resume_tailor.docx_io import (
    document_blocks,
    extract_markdown,
    nonempty_paragraphs,
    set_paragraph_text,
    verify_written_docx,
    write_tailored_docx,
)

HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(url, HYPERLINK_REL, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


@pytest.fixture
def sample_docx(tmp_path):
    """A synthetic resume shaped like the real one. No production content."""
    document = Document()

    name = document.add_paragraph()
    name.add_run("        ")           # centring by leading spaces
    name.add_run("Alex Placeholder (SAMPLE)").bold = True

    contact = document.add_paragraph()
    _add_hyperlink(contact, "not-a-real-person@example.invalid", "mailto:x@example.invalid")
    contact.add_run(" | +1-555-0100")

    document.add_paragraph("PROFESSIONAL SUMMARY:")
    document.add_paragraph("Engineer with 9+ years of widget experience.", style="List Bullet")

    document.add_paragraph("TECHNICAL SKILLS:")
    skills = document.add_paragraph()
    skills.add_run("Languages: ").bold = True     # bold label, regular tail
    skills.add_run("Python, SQL, Bash")

    document.add_paragraph("PROFESSIONAL EXPERIENCE:")
    document.add_paragraph("Northstar Fictional Bank, Testland | Jan 2024 – Present")
    document.add_paragraph("Senior Widget Engineer (SAMPLE)")
    document.add_paragraph("Built widgets.", style="List Bullet")
    document.add_paragraph("Maintained widget pipelines.", style="List Bullet")

    path = tmp_path / "sample.docx"
    document.save(str(path))
    return path


def _texts(path):
    return [p.text for p in nonempty_paragraphs(Document(str(path)))]


def test_unchanged_markdown_produces_an_unchanged_document(sample_docx, tmp_path):
    """The regression that matters most: a no-op edit must be a genuine no-op."""
    markdown = extract_markdown(sample_docx)
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    assert _texts(dest) == _texts(sample_docx)


def test_bullet_edit_lands_in_the_document(sample_docx, tmp_path):
    markdown = extract_markdown(sample_docx).replace(
        "- Built widgets.", "- Built production widgets."
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    assert "Built production widgets." in _texts(dest)


def test_editing_a_skills_line_does_not_bold_the_whole_line(sample_docx, tmp_path):
    """Collapsing runs into runs[0] used to make every edited skills line bold."""
    markdown = extract_markdown(sample_docx).replace(
        "Languages: Python, SQL, Bash", "Languages: Python, SQL, Bash, Go"
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    paragraph = next(
        p for p in nonempty_paragraphs(Document(str(dest))) if p.text.startswith("Languages:")
    )
    assert "Go" in paragraph.text
    bolded = [r.text for r in paragraph.runs if r.text.strip() and r.bold]
    assert bolded == ["Languages: "], "only the label should stay bold"


def test_a_new_line_is_inserted_rather_than_dropped(sample_docx, tmp_path):
    """map_paragraphs could only overwrite 1:1, so additions vanished silently."""
    markdown = extract_markdown(sample_docx).replace(
        "Languages: Python, SQL, Bash",
        "Languages: Python, SQL, Bash\nPlatform: Kubernetes, Docker",
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    texts = _texts(dest)
    assert any(t.startswith("Platform: Kubernetes") for t in texts)
    assert len(texts) == len(_texts(sample_docx)) + 1


def test_a_removed_line_is_deleted_from_the_document(sample_docx, tmp_path):
    markdown = extract_markdown(sample_docx).replace(
        "- Maintained widget pipelines.\n", ""
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    texts = _texts(dest)
    assert not any("Maintained widget pipelines" in t for t in texts)
    assert len(texts) == len(_texts(sample_docx)) - 1


def test_hyperlink_paragraph_is_never_rewritten(sample_docx, tmp_path):
    """Hyperlink text sits outside paragraph.runs, so a rewrite duplicated it."""
    document = Document(str(sample_docx))
    contact = nonempty_paragraphs(document)[1]
    before = contact.text
    changed = set_paragraph_text(contact, "someone-else@example.invalid | +1-555-9999")
    assert changed is False
    assert contact.text == before
    assert contact.text.count("not-a-real-person@example.invalid") == 1


def test_leading_whitespace_centring_survives(sample_docx, tmp_path):
    markdown = extract_markdown(sample_docx)
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    assert _texts(dest)[0].startswith("        ")


def test_bullets_cannot_migrate_between_sections(sample_docx, tmp_path):
    """Ordered alignment: an edit near the end must not rewrite an early bullet."""
    markdown = extract_markdown(sample_docx).replace(
        "- Maintained widget pipelines.", "- Maintained and scaled widget pipelines."
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    blocks, _ = document_blocks(Document(str(dest)))
    summary = next(b for b in blocks if "9+ years" in b.text)
    assert summary.text == "Engineer with 9+ years of widget experience."


def test_curly_quote_normalisation_is_not_treated_as_an_edit(sample_docx, tmp_path):
    """Models swap ’ for '. That is typography, not tailoring — leave the run alone."""
    document = Document(str(sample_docx))
    paragraph = nonempty_paragraphs(document)[3]
    paragraph.runs[0].text = "Engineer’s widget experience."
    before = paragraph.text
    assert set_paragraph_text(paragraph, "Engineer's widget experience.") is False
    assert paragraph.text == before


def test_an_inserted_skills_line_does_not_inherit_the_bold_label(sample_docx, tmp_path):
    """A cloned paragraph used to put the whole line in the template's bold run."""
    markdown = extract_markdown(sample_docx).replace(
        "Languages: Python, SQL, Bash",
        "Languages: Python, SQL, Bash\nPlatform: Kubernetes, Docker, Terraform",
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    inserted = next(
        p for p in nonempty_paragraphs(Document(str(dest)))
        if p.text.startswith("Platform:")
    )
    bolded = [r.text for r in inserted.runs if r.text.strip() and r.bold]
    assert bolded == ["Platform: "], f"only the label should be bold, got {bolded}"


def test_consecutive_inserts_keep_their_order(sample_docx, tmp_path):
    """Every clone used to be placed after the same anchor, reversing the run."""
    markdown = extract_markdown(sample_docx).replace(
        "- Built widgets.",
        "- Built widgets.\n- Added ONE\n- Added TWO\n- Added THREE",
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    added = [t for t in _texts(dest) if t.startswith("Added ")]
    assert added == ["Added ONE", "Added TWO", "Added THREE"]


def test_a_kind_changing_replacement_loses_nothing(sample_docx, tmp_path):
    """Insert anchored on an already-deleted paragraph dropped both lines."""
    markdown = extract_markdown(sample_docx).replace(
        "Languages: Python, SQL, Bash", "- Shipped a widget pipeline end to end"
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    texts = _texts(dest)
    assert any("Shipped a widget pipeline" in t for t in texts)
    assert not any(t.startswith("Languages:") for t in texts)
    assert len(texts) == len(_texts(sample_docx))


def test_verification_passes_for_a_faithful_write(sample_docx, tmp_path):
    markdown = extract_markdown(sample_docx).replace(
        "- Built widgets.", "- Built production widgets.\n- Added a new bullet"
    )
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)
    assert verify_written_docx(dest, markdown) == []


def test_verification_catches_a_line_lost_on_the_way_to_disk(sample_docx, tmp_path):
    """Nothing else compares the delivered file to what the guardrails approved."""
    markdown = extract_markdown(sample_docx)
    dest = tmp_path / "out.docx"
    write_tailored_docx(sample_docx, markdown, dest)

    document = Document(str(dest))
    paragraphs = nonempty_paragraphs(document)
    paragraphs[3]._p.getparent().remove(paragraphs[3]._p)
    document.save(str(dest))

    problems = verify_written_docx(dest, markdown)
    assert problems and "missing" in problems[0].lower()


def test_edit_starting_inside_the_bold_label_does_not_bold_the_line(sample_docx):
    """keep < len(runs[0]) used to let the bold label run absorb the whole line."""
    document = Document(str(sample_docx))
    paragraph = next(
        p for p in nonempty_paragraphs(document) if p.text.startswith("Languages:")
    )
    set_paragraph_text(paragraph, "Core Languages: Python, Go, SQL")
    bolded = [r.text for r in paragraph.runs if r.text.strip() and r.bold]
    assert bolded == ["Core Languages: "]
